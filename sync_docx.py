import sys
import os
import re

def ensure_dependencies():
    try:
        import docx
    except ImportError:
        print("python-docx not found. Trying to install...")
        import subprocess
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
            import docx
        except Exception as e:
            print(f"Error installing python-docx: {e}")
            sys.exit(1)
    return docx

docx_module = ensure_dependencies()
from docx import Document
from docx.shared import Inches, Pt, RGBColor

def add_formatted_runs(paragraph, text):
    # Split by bold tags **
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
        else:
            # Check for simple markdown links [text](url)
            link_parts = re.split(r'(\[.*?\]\(.*?\))', part)
            for l_part in link_parts:
                if l_part.startswith('[') and ']' in l_part and '(' in l_part and l_part.endswith(')'):
                    m = re.match(r'^\[(.*?)\]\((.*?)\)$', l_part)
                    if m:
                        link_text = m.group(1)
                        link_url = m.group(2)
                        run = paragraph.add_run(f"{link_text} ({link_url})")
                        run.font.underline = True
                        run.font.color.rgb = RGBColor(0, 51, 153) # Dark blue for links
                    else:
                        paragraph.add_run(l_part)
                else:
                    paragraph.add_run(l_part)

def md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} does not exist")
        return False
    
    print(f"Reading Markdown from {md_path}...")
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    doc = Document()
    
    # Page Setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Set default paragraph style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    
    for line in lines:
        stripped = line.strip()
        
        # 1. Empty lines
        if not stripped:
            continue
            
        # 2. Horizontal rule
        if stripped == '---' or stripped == '***':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('—' * 50)
            run.font.color.rgb = RGBColor(180, 180, 180)
            continue
            
        # 3. Headings
        h_match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2)
            p = doc.add_heading(level=level)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            # Apply heading font size & formatting
            run = p.runs[0] if p.runs else p.add_run()
            # docx default heading style might not use the font we set, let's format it
            run.text = ""
            add_formatted_runs(p, text)
            for r in p.runs:
                r.font.name = 'Microsoft YaHei'
                r.bold = True
                if level == 1:
                    r.font.size = Pt(18)
                    r.font.color.rgb = RGBColor(31, 78, 121)
                elif level == 2:
                    r.font.size = Pt(14)
                    r.font.color.rgb = RGBColor(46, 116, 181)
                else:
                    r.font.size = Pt(12)
                    r.font.color.rgb = RGBColor(89, 89, 89)
            continue
            
        # 4. Bullet list items (* or -)
        bullet_match = re.match(r'^[\*\-]\s+(.*)$', stripped)
        if bullet_match:
            text = bullet_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            add_formatted_runs(p, text)
            # Make sure list runs use Microsoft YaHei
            for r in p.runs:
                r.font.name = 'Microsoft YaHei'
            continue
            
        # 5. Numbered list items (e.g. 1. or 2.)
        num_match = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if num_match:
            text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            add_formatted_runs(p, text)
            for r in p.runs:
                r.font.name = 'Microsoft YaHei'
            continue
            
        # 6. Indented list items (sub-bullets with leading spaces)
        sub_bullet_match = re.match(r'^\s{2,}[\*\-]\s+(.*)$', line)
        if sub_bullet_match:
            text = sub_bullet_match.group(1)
            p = doc.add_paragraph(style='List Bullet 2')
            p.paragraph_format.space_after = Pt(2)
            add_formatted_runs(p, text)
            for r in p.runs:
                r.font.name = 'Microsoft YaHei'
            continue
            
        # 7. Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        add_formatted_runs(p, stripped)
        for r in p.runs:
            r.font.name = 'Microsoft YaHei'
            
    doc.save(docx_path)
    print(f"Successfully saved docx to {docx_path}")
    return True

if __name__ == '__main__':
    md_file = r'D:\heming\shike\文档\04_运营与冷启动\04_小红书公开构建素材(第三期).md'
    docx_file = r'D:\heming\shike\文档\04_运营与冷启动\04_小红书公开构建素材(第三期).docx'
    
    if len(sys.argv) > 1:
        md_file = sys.argv[1]
    if len(sys.argv) > 2:
        docx_file = sys.argv[2]
        
    md_to_docx(md_file, docx_file)
